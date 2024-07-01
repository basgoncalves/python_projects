# USAGE:
# run the code and you will be asked to select the Volume folder 
# The assumption is that you have access to UniVienna Shared Drive "Z:\iacss\IACSS_IJCSS\IJCSS\Volumes"
# If you don't have access to it ask the IT office

import install_needed_pkg
import subprocess 
import aspose.words as aw
import os
from tkinter import Tk
from tkinter.filedialog import askopenfilename, askdirectory
from docx2pdf import convert
import shutil
import textract
import docx

from tkinter import filedialog

def check_volumes_path():
    volumes_path = r'Z:\iacss\IACSS_IJCSS\IJCSS\Volumes'
    if not os.path.exists(volumes_path):
        volumes_path = filedialog.askdirectory(title="Select path of volumes folder")
    print("Volumes path:", volumes_path)
    return volumes_path

def loop_through_doi_txt():    
    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doi_to_check.txt')
    with open(file_path, 'r') as file:
        for line in file:
            print(line)          

def please_close_word():
    import pyautogui
    import time

    # Check if any Word documents are open
    while pyautogui.getWindowsWithTitle("Word"):
        # Ask the user to save and close all documents
        result = pyautogui.confirm(
            "Some Word documents are still open. Do you want to save and close them?",
            buttons=["Yes", "No", "Cancel"]
        )
        
        if result == "Yes":
            word_windows = pyautogui.getWindowsWithTitle("Word")
            # Press Alt+F4 to close all Word documents
            for window in word_windows:
                window.close()
                pyautogui.alert("All Word documents are closed.")
        elif result == "Cancel":
            return False
           
    # if function reaches this point, it means that all Word documents are closed
    return True

def extract_text_between_abstract_and_keywords(doc_path):
    doc = docx.Document(doc_path)
    start_extraction = False
    extracted_text = []

    for paragraph in doc.paragraphs:
        if "abstract" in paragraph.text.lower():
            start_extraction = True
            continue
        if "keywords" in paragraph.text.lower() and start_extraction:
            break
        if start_extraction:
            extracted_text.append(paragraph.text)

    return "\n".join(extracted_text)

def extract_first_two_lines(doc_path):
    doc = docx.Document(doc_path)
    first_line = []

    for i, paragraph in enumerate(doc.paragraphs):
        first_line.append(paragraph.text)
        if i >= 2:
            break

    return first_line

def modify_document(input_doc_path, output_doc_path, new_text,line_starts_with):
    doc = docx.Document(input_doc_path)
    
    # Find the paragraph that starts with line_starts_with and delete everything after it
    for paragraph in doc.paragraphs:
        
        if paragraph.text.startswith(line_starts_with):
            paragraph.clear()
            run = paragraph.add_run(line_starts_with)
            bold_format = run.font
            bold_format.bold = True
            paragraph.add_run(new_text)
            break

    # Save the modified document
    doc.save(output_doc_path)

def add_to_content_specification(volumes_path,DOI,output_document_path):
    
    # output_document_path = os.path.join(submit_folder, 'Content Specification_{}.docx'.format(DOI))
    word_path = os.path.join(volumes_path,'3-Uploads/1-Word')
    main_document = os.path.join(word_path, '{}.docx'.format(DOI))
    template_content_specification = os.path.join(volumes_path, 'Template_Content Specification.docx')
    # output_document_path = r'Z:\iacss\IACSS_IJCSS\IJCSS\Volumes\Vol222023Ed2\3-Uploads\10516-Volume22-Issue2c\Content Specification.docx'

    # filename and DOI
    filename =  DOI + '.pdf'
    modify_document(template_content_specification, output_document_path, filename,"Filename: ")
    modify_document(output_document_path, output_document_path, DOI,"DOI: ")

    # abstract
    extracted_text = extract_text_between_abstract_and_keywords(main_document)
    modify_document(output_document_path, output_document_path, extracted_text,"Abstract: ")

    # title and authors
    first_lines_text = extract_first_two_lines(main_document)
    modify_document(output_document_path, output_document_path, first_lines_text[0],"Title: ")
    modify_document(output_document_path, output_document_path, first_lines_text[1],"Author(s): ")
    modify_document(output_document_path, output_document_path, first_lines_text[2],"Affiliation(s): ")

def create_final_pdf(volumes_path,re_convert_files_that_already_exist):
    
    # get volume and issue from folder name (must be in the format 'Vol222023Ed1')
    volume = volumes_path.split(r'/')[-1].split('Vol')[1][0:2]   
    issue = volumes_path.split(r'/')[-1].split('Ed')[1][0:1]

    ghostscript_path = r"C:\Program Files\gs\gs10.02.1\bin\gswin64c.exe"  # Update with your actual path

    # create final submission folder (will be used to zip in the end)
    submit_folder = os.path.join(volumes_path, r'3-Uploads\10516-Volume{}-Issue{}'.format(volume,issue), '')

    if not os.path.isdir(submit_folder):
        os.mkdir(submit_folder)
    
    word_path = os.path.join(volumes_path, r'3-Uploads\1-Word','')
    pdf_path = os.path.join(volumes_path, r'3-Uploads\2-PDF_ausWord','')
    ps_path = os.path.join(volumes_path,r'3-Uploads\3-PostScript','')
    final_path = os.path.join(volumes_path, r'3-Uploads\4-PDF_ausPostScript','')
    
    # create all the subfiles (PDF_ausWord, 3-PostScript, and PDF_ausPostScript for each of the word docs)
    for path, subdirs, files in os.walk(word_path):
        for name in files:
            # For each file, ensure it is a .docx file before adding it to the list
            try: 
                text = textract.process(os.path.join(path, name))
            except: 
                print('not able to open ' + os.path.join(path, name))
                continue
            
            if os.path.splitext(os.path.join(path, name))[1] == ".docx":
                word_file_path = os.path.join(path, name)
                
                # assign the names of each file format (pdf, ps, final_pdf)
                file_path, file_extension = os.path.splitext(word_file_path)
                pdf_file_path = os.path.join(pdf_path,name.replace(file_extension,'.pdf'))
                ps_file_path = os.path.join(ps_path,name.replace(file_extension,'.eps'))
                pdf_file_path_final = os.path.join(final_path,name.replace(file_extension,'.pdf'))
                submit_file_path_final = os.path.join(submit_folder,name.replace(file_extension,'.pdf'))
                        
                
                if not os.path.isfile(submit_file_path_final) or re_convert_files_that_already_exist == True:
                    
                    convert(word_file_path, pdf_file_path)                          # convert to PDF: "1-Word"

                    if not os.path.isfile(ghostscript_path):
                        print('\n \n \n \n ERROR: Ghostscript path not found in {}'.format(ghostscript_path))
                        print('Please ensure Ghostscrpt is installed or update the path in the script.')
                        exit()

                    # convert to post script: "2-PDF_ausWord"
                    try:
                        try:
                            subprocess.call(['pdf2ps', pdf_file_path, ps_file_path])        # convert to post script: "2-PDF_ausWord"
                        except:
                            subprocess.call([ghostscript_path, '-sDEVICE=ps2write', '-o', ps_file_path, pdf_file_path])
                    except  Exception as e:
                        print('Error converting to post script: {}'.format(e))
                        exit()

                    # convert to pdf again: "3-PostScript"
                    try:
                        try:
                            subprocess.call(['ps2pdf', ps_file_path, pdf_file_path_final])  # convert to pdf again: "3-PostScript"
                        except:
                            subprocess.call([ghostscript_path, '-sDEVICE=pdfwrite', '-o', pdf_file_path_final, ps_file_path])
                    except  Exception as e:
                        print('Error converting to pdf again: {}'.format(e))
                        exit()
                    
                    # copy pdf to the final location: "4-PDF_ausPostScript"
                    shutil.copyfile(pdf_file_path_final, submit_file_path_final)    
                    
                    # create content specification
                    DOI = os.path.splitext(os.path.basename(word_file_path))[0]
                    content_spec_path = os.path.join(submit_folder, 'Content Specification_{}.docx'.format(DOI))
                    add_to_content_specification(volumes_path,DOI,content_spec_path)
                    
                    # copy files to a single folder in case there is a need to uplad individually
                    las_digits_DOI = DOI.split("-")[-1]
                    submit_folder_per_paper = os.path.dirname(submit_folder) + '_' + las_digits_DOI
                    
                    # create folder if it does not exist
                    if not os.path.isdir(submit_folder_per_paper):
                        os.makedirs(submit_folder_per_paper)
                        
                    # copy content spec
                    document_name = os.path.basename(content_spec_path)
                    destination_path = os.path.join(submit_folder_per_paper,'{}'.format(document_name))
                    shutil.copyfile(content_spec_path, destination_path)
                    
                    # copy pdf 
                    document_name = os.path.basename(submit_file_path_final)
                    destination_path = os.path.join(submit_folder_per_paper,'{}'.format(document_name))
                    shutil.copyfile(submit_file_path_final, destination_path)


if __name__ == '__main__':    
    out = please_close_word()
    if out == False:
        exit()
    # if you want old files to be converted again, change to "True"
    re_convert_files_that_already_exist = False

    # User select the folder of the volume you want to convert folders 
    volumes_path = check_volumes_path()
    edition_path = askdirectory(initialdir=r'Z:\iacss\IACSS_IJCSS\IJCSS\Volumes')

    os.startfile(os.path.join(os.path.dirname (edition_path), r'IJCSS - Progress_all.xlsx'))

    create_final_pdf(edition_path,re_convert_files_that_already_exist)


# end